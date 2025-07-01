import tkinter as tk
from tkinter import ttk, filedialog, messagebox, Toplevel, Listbox, END, Menu
import json
import os
import re
import docx
import sv_ttk
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- 全局配置 ---
CONFIG_FILE = 'gov_doc_format_config.json'

# --- 格式化引擎 (稳定版) ---
class GovDocFormatter:
    FONT_XBS = '方正小标宋简体'
    FONT_FS = '仿宋_GB2312'
    FONT_HT = '黑体'
    FONT_KT = '楷体_GB2312'
    FONT_SONG = '宋体'
    SIZE_2 = Pt(22)
    SIZE_3 = Pt(16)
    SIZE_4 = Pt(12)

    def set_font_style(self, run, font_name, font_size, bold=False, color_rgb=None):
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        if color_rgb:
            run.font.color.rgb = RGBColor(color_rgb[0], color_rgb[1], color_rgb[2])
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def set_paragraph_format(self, p, alignment=None, line_spacing=None, space_before=Pt(0), space_after=Pt(0), first_line_indent=None, right_indent=None):
        if alignment is not None: p.alignment = alignment
        if line_spacing is not None:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = line_spacing
        p.paragraph_format.space_before = space_before
        p.paragraph_format.space_after = space_after
        if first_line_indent is not None: p.paragraph_format.first_line_indent = first_line_indent
        if right_indent is not None: p.paragraph_format.right_indent = right_indent

    def add_separator(self, doc, thickness='single', size=4, color='FF0000', space_before=Mm(4)):
        p = doc.add_paragraph()
        self.set_paragraph_format(p, space_before=space_before, space_after=Pt(0))
        pPr = p._element.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom_bdr = OxmlElement('w:bottom')
        bottom_bdr.set(qn('w:val'), thickness)
        bottom_bdr.set(qn('w:sz'), str(size))
        bottom_bdr.set(qn('w:color'), color)
        p_bdr.append(bottom_bdr)
        pPr.append(p_bdr)

    def add_page_number(self, doc):
        doc.settings.odd_and_even_pages_header_footer = True
        for section in doc.sections:
            self._create_footer_page_number(section.footer, WD_ALIGN_PARAGRAPH.RIGHT)
            self._create_footer_page_number(section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT)

    def _create_footer_page_number(self, footer, alignment):
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = alignment
        p.clear()
        run = p.add_run('— '); self.set_font_style(run, self.FONT_SONG, self.SIZE_4)
        fld_char_begin = OxmlElement('w:fldChar'); fld_char_begin.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText'); instr_text.set(qn('xml:space'), 'preserve'); instr_text.text = 'PAGE'
        fld_char_end = OxmlElement('w:fldChar'); fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_begin); run._r.append(instr_text); run._r.append(fld_char_end)
        run = p.add_run(' —'); self.set_font_style(run, self.FONT_SONG, self.SIZE_4)

    def process(self, data, input_path):
        doc = docx.Document()
        section = doc.sections[0]
        section.top_margin = Mm(37); section.bottom_margin = Mm(35); section.left_margin = Mm(28); section.right_margin = Mm(26)

        if data.get('add_page_number'): self.add_page_number(doc)

        # --- 版头 ---
        if data.get('copy_number'):
            p = doc.add_paragraph(); self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(data['copy_number']); self.set_font_style(run, self.FONT_FS, self.SIZE_3)
        if data.get('security_level'):
            p = doc.add_paragraph(); self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(data['security_level']); self.set_font_style(run, self.FONT_HT, self.SIZE_3)
        if data.get('urgency'):
            p = doc.add_paragraph(); self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(data['urgency']); self.set_font_style(run, self.FONT_HT, self.SIZE_3)

        if data.get('issuing_authority_logo'):
            p = doc.add_paragraph()
            self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Mm(35) - self.SIZE_2)
            run = p.add_run(data['issuing_authority_logo']); self.set_font_style(run, self.FONT_XBS, self.SIZE_2, color_rgb=(255, 0, 0))

        if data.get('doc_number') or data.get('signatory'):
            p = doc.add_paragraph()
            self.set_paragraph_format(p, space_before=self.SIZE_3 * 2)
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Mm(156), WD_TAB_ALIGNMENT.RIGHT)
            if data.get('doc_number'):
                run_num = p.add_run(data['doc_number']); self.set_font_style(run_num, self.FONT_FS, self.SIZE_3)
            if data.get('signatory'):
                p.add_run('	')
                run_sig_label = p.add_run('签发人：'); self.set_font_style(run_sig_label, self.FONT_FS, self.SIZE_3)
                run_sig_name = p.add_run(data['signatory']); self.set_font_style(run_sig_name, self.FONT_KT, self.SIZE_3)

        if data.get('add_red_separator'): self.add_separator(doc)

        # --- 主体 ---
        title_source = data.get('title_option', 'auto')
        main_title = ""
        if title_source == 'manual':
            main_title = data.get('main_title_manual', '')
        elif input_path:
            try:
                if input_path.lower().endswith('.txt'):
                    with open(input_path, 'r', encoding='utf-8') as f: main_title = f.readline().strip()
                elif input_path.lower().endswith('.docx'):
                    source_doc = docx.Document(input_path)
                    if source_doc.paragraphs: main_title = source_doc.paragraphs[0].text.strip()
            except Exception: pass
        
        if main_title:
            p = doc.add_paragraph(); self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=self.SIZE_3 * 2)
            run = p.add_run(main_title); self.set_font_style(run, self.FONT_XBS, self.SIZE_2)

        if data.get('main_recipient'):
            p = doc.add_paragraph(); self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=self.SIZE_3)
            run = p.add_run(data['main_recipient'] + '：'); self.set_font_style(run, self.FONT_FS, self.SIZE_3)

        doc.add_paragraph()
        self.process_body(doc, input_path, main_title if title_source == 'auto' else None)

        # --- 文末要素 ---
        if data.get('attachment_note'):
            p = doc.add_paragraph(); self.set_paragraph_format(p, line_spacing=Pt(28), first_line_indent=Pt(self.SIZE_3.pt * 2), space_before=self.SIZE_3)
            run = p.add_run('附件：' + data['attachment_note']); self.set_font_style(run, self.FONT_FS, self.SIZE_3)

        if data.get('issuing_authority_signature') and data.get('doc_date'):
            indent_chars = 4 if data.get('is_stamped') else 2
            p = doc.add_paragraph(); self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, right_indent=Pt(self.SIZE_3.pt * indent_chars), space_before=self.SIZE_3 * 2)
            run_sig = p.add_run(data['issuing_authority_signature']); self.set_font_style(run_sig, self.FONT_FS, self.SIZE_3)
            p.add_run('\n')
            run_date = p.add_run(data['doc_date']); self.set_font_style(run_date, self.FONT_FS, self.SIZE_3)

        if data.get('addendum'):
            p = doc.add_paragraph(); self.set_paragraph_format(p, line_spacing=Pt(28), first_line_indent=Pt(self.SIZE_3.pt * 2))
            run_text = "（" + data.get('addendum', '') + "）"
            run = p.add_run(run_text); self.set_font_style(run, self.FONT_FS, self.SIZE_3)

        # --- 版记 ---
        if data.get('cc_list') or data.get('printing_info'):
            self.add_separator(doc, thickness='double', size=6, color='000000', space_before=self.SIZE_3)
            if data.get('cc_list'):
                p = doc.add_paragraph(); self.set_paragraph_format(p, line_spacing=Pt(28))
                run = p.add_run("抄送：" + data.get('cc_list','')); self.set_font_style(run, self.FONT_FS, self.SIZE_4)
            if data.get('printing_info'):
                p = doc.add_paragraph(); self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=Pt(28))
                run = p.add_run(data['printing_info']); self.set_font_style(run, self.FONT_FS, self.SIZE_4)
            self.add_separator(doc, thickness='single', size=6, color='000000', space_before=Pt(0))

        output_path = os.path.splitext(input_path)[0] + "_formatted.docx"
        doc.save(output_path)
        return output_path

    def process_body(self, doc, input_path, auto_detected_title):
        is_docx = input_path.lower().endswith('.docx')
        if is_docx:
            source_doc = docx.Document(input_path)
            elements = source_doc.element.body
        else:
            with open(input_path, 'r', encoding='utf-8') as f:
                elements = f.readlines()

        for i, element in enumerate(elements):
            text_to_process, is_table = "", False
            if is_docx:
                if element.tag.endswith('p'): text_to_process = docx.text.paragraph.Paragraph(element, doc).text.strip()
                elif element.tag.endswith('tbl'): is_table = True
            else: text_to_process = element.strip()

            if is_table:
                self._format_table(doc, docx.table.Table(element, doc)); continue
            if not text_to_process: continue
            if i == 0 and text_to_process == auto_detected_title: continue
            self._format_paragraph(doc, text_to_process)

    def _format_paragraph(self, doc, text):
        p = doc.add_paragraph()
        self.set_paragraph_format(p, line_spacing=Pt(28))
        if re.match(r'^一、|^二、|^三、', text): run = p.add_run(text); self.set_font_style(run, self.FONT_HT, self.SIZE_3)
        elif re.match(r'^\uff08[一二三四五]+\uff09', text): run = p.add_run(text); self.set_font_style(run, self.FONT_KT, self.SIZE_3, bold=True)
        elif re.match(r'^\d+\.', text): run = p.add_run(text); self.set_font_style(run, self.FONT_FS, self.SIZE_3, bold=True)
        else:
            self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=Pt(28), first_line_indent=Pt(self.SIZE_3.pt * 2))
            run = p.add_run(text); self.set_font_style(run, self.FONT_FS, self.SIZE_3)

    def _format_table(self, doc, source_table):
        new_table = doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
        new_table.style = 'Table Grid'
        new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        new_table.autofit = True
        for i, row in enumerate(source_table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.cell(i, j)
                new_cell.text = ""
                for para in cell.paragraphs:
                    new_para = new_cell.add_paragraph(para.text)
                    self.set_paragraph_format(new_para, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=Pt(28))
                    for run in new_para.runs:
                        self.set_font_style(run, self.FONT_FS, self.SIZE_3)

# --- GUI and App Logic ---
class ConfigManager:
    def __init__(self, filename=CONFIG_FILE):
        self.filename = filename; self.data = self.load()
    def load(self):
        try:
            with open(self.filename, 'r', encoding='utf-8') as f: return json.load(f)
        except (FileNotFoundError, json.JSONDecodeError): return {}
    def save(self):
        with open(self.filename, 'w', encoding='utf-8') as f: json.dump(self.data, f, ensure_ascii=False, indent=4)
    def get(self, key, default=None): return self.data.get(key, default or [])
    def set(self, key, value): self.data[key] = value; self.save()

class ManagementDialog(Toplevel):
    def __init__(self, parent, title, key, config_manager, on_close_callback):
        super().__init__(parent)
        self.title(f"管理 {title}"); self.key = key; self.config_manager = config_manager
        self.on_close_callback = on_close_callback
        self.protocol("WM_DELETE_WINDOW", self.close_dialog)
        self.transient(parent); self.grab_set()

        self.listbox = Listbox(self, width=50, height=10, font=("Segoe UI", 10)); self.listbox.pack(padx=15, pady=15, fill=tk.BOTH, expand=True)
        for item in self.config_manager.get(self.key): self.listbox.insert(END, item)
        
        btn_frame = ttk.Frame(self); btn_frame.pack(fill=tk.X, padx=15, pady=(0, 15))
        del_btn = ttk.Button(btn_frame, text="删除选中", command=self.delete_item); del_btn.pack(side=tk.RIGHT)

    def delete_item(self):
        selected_indices = self.listbox.curselection()
        if not selected_indices: return
        for i in reversed(selected_indices):
            self.listbox.delete(i)
        self.save_changes()

    def save_changes(self):
        items = list(self.listbox.get(0, END))
        self.config_manager.set(self.key, items)

    def close_dialog(self):
        self.on_close_callback()
        self.destroy()

class ManagedField(ttk.Frame):
    def __init__(self, parent, key, label, config_manager, controls_dict):
        super().__init__(parent)
        self.key = key
        self.config_manager = config_manager
        self.controls_dict = controls_dict

        self.pack(fill=tk.X, pady=10, anchor='n')
        ttk.Label(self, text=label, font=("Segoe UI", 10, "bold")).pack(anchor=tk.W, pady=(0, 5))
        
        input_frame = ttk.Frame(self)
        input_frame.pack(fill=tk.X)

        self.combo = ttk.Combobox(input_frame, font=("Segoe UI", 10))
        self.combo.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.controls_dict[key] = self.combo
        self.refresh_values()

        add_button = ttk.Button(input_frame, text="+ ", width=3, command=self.add_item)
        add_button.pack(side=tk.LEFT, padx=(5, 5))

        manage_button = ttk.Button(input_frame, text="⚙️", width=3, command=self.open_management_dialog)
        manage_button.pack(side=tk.LEFT)

    def add_item(self):
        new_item = self.combo.get()
        if not new_item: return
        current_items = self.config_manager.get(self.key, [])
        if new_item not in current_items:
            current_items.append(new_item)
            self.config_manager.set(self.key, current_items)
            self.refresh_values()
            self.combo.set(new_item)

    def open_management_dialog(self):
        ManagementDialog(self, self.key, self.key, self.config_manager, self.refresh_values)

    def refresh_values(self):
        self.combo['values'] = self.config_manager.get(self.key, [])

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("公文智能排版工具"); self.geometry("800x850")
        self.config_manager = ConfigManager()
        self.formatter = GovDocFormatter()
        sv_ttk.set_theme("dark")
        self.create_widgets()

    def create_widgets(self):
        main_frame = ttk.Frame(self, padding=25); main_frame.pack(fill=tk.BOTH, expand=True)
        self.controls = {}

        file_card = ttk.LabelFrame(main_frame, text=" 源文件 ", padding=20)
        file_card.pack(fill=tk.X, pady=(0, 15))
        self.file_path_var = tk.StringVar(value="尚未选择文件...")
        select_button = ttk.Button(file_card, text="选择 .txt 或 .docx 文件", command=self.select_file)
        select_button.pack(side=tk.LEFT, padx=(0, 15))
        ttk.Label(file_card, textvariable=self.file_path_var).pack(side=tk.LEFT)

        notebook = ttk.Notebook(main_frame); notebook.pack(fill=tk.BOTH, expand=True, pady=10)
        tab1 = ttk.Frame(notebook, padding=20); tab2 = ttk.Frame(notebook, padding=20); tab3 = ttk.Frame(notebook, padding=20)
        notebook.add(tab1, text='版头要素'); notebook.add(tab2, text='主体与文末'); notebook.add(tab3, text='版记与选项')

        # --- Tabs Content ---
        ManagedField(tab1, "copy_number", "份号", self.config_manager, self.controls)
        ManagedField(tab1, "security_level", "密级和保密期限", self.config_manager, self.controls)
        ManagedField(tab1, "urgency", "紧急程度", self.config_manager, self.controls)
        ManagedField(tab1, "issuing_authority_logo", "发文机关标志", self.config_manager, self.controls)
        ManagedField(tab1, "doc_number", "发文字号", self.config_manager, self.controls)
        ManagedField(tab1, "signatory", "签发人", self.config_manager, self.controls)

        self.create_title_options(tab2)
        ManagedField(tab2, "main_recipient", "主送机关", self.config_manager, self.controls)
        ManagedField(tab2, "issuing_authority_signature", "发文机关署名", self.config_manager, self.controls)
        ManagedField(tab2, "doc_date", "成文日期", self.config_manager, self.controls)
        ManagedField(tab2, "attachment_note", "附件说明", self.config_manager, self.controls)
        ManagedField(tab2, "addendum", "附注", self.config_manager, self.controls)

        ManagedField(tab3, "cc_list", "抄送机关", self.config_manager, self.controls)
        ManagedField(tab3, "printing_info", "印发机关和印发日期", self.config_manager, self.controls)
        
        options_card = ttk.LabelFrame(tab3, text=" 格式选项 ", padding=20)
        options_card.pack(fill=tk.X, pady=25)
        self.controls['add_red_separator'] = tk.BooleanVar(value=False)
        ttk.Checkbutton(options_card, text="生成红色分隔线", variable=self.controls['add_red_separator']).pack(side=tk.LEFT, padx=15)
        self.controls['is_stamped'] = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_card, text="是否加盖印章", variable=self.controls['is_stamped']).pack(side=tk.LEFT, padx=15)
        self.controls['add_page_number'] = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_card, text="生成页码", variable=self.controls['add_page_number']).pack(side=tk.LEFT, padx=15)

        generate_button = ttk.Button(main_frame, text="生成格式化Word文档", command=self.generate_document, style='Accent.TButton')
        generate_button.pack(fill=tk.X, ipady=8, pady=15)

    def create_title_options(self, parent):
        container = ttk.LabelFrame(parent, text=" 公文标题 ", padding=20)
        container.pack(fill=tk.X, pady=10)
        
        self.controls['title_option'] = tk.StringVar(value="auto")
        
        auto_rb = ttk.Radiobutton(container, text="自动获取 (源文件第一行)", variable=self.controls['title_option'], value="auto", command=self.toggle_manual_title)
        auto_rb.pack(anchor=tk.W)
        
        manual_frame = ttk.Frame(container)
        manual_frame.pack(fill=tk.X, anchor=tk.W, pady=(5,0))
        manual_rb = ttk.Radiobutton(manual_frame, text="手动指定", variable=self.controls['title_option'], value="manual", command=self.toggle_manual_title)
        manual_rb.pack(side=tk.LEFT, anchor=tk.W)
        
        self.controls['main_title_manual'] = ttk.Entry(manual_frame, state=tk.DISABLED, font=("Segoe UI", 10))
        self.controls['main_title_manual'].pack(side=tk.LEFT, fill=tk.X, expand=True, padx=10)

    def toggle_manual_title(self):
        state = tk.NORMAL if self.controls['title_option'].get() == 'manual' else tk.DISABLED
        self.controls['main_title_manual'].config(state=state)

    def select_file(self):
        path = filedialog.askopenfilename(filetypes=[("All supported", "*.txt *.docx"), ("Text", "*.txt"), ("Word", "*.docx")])
        if path: self.file_path_var.set(path)

    def generate_document(self):
        input_path = self.file_path_var.get()
        if "尚未选择" in input_path: messagebox.showwarning("警告", "请先选择一个文件！"); return
        gui_data = {key: (var.get() if hasattr(var, 'get') else var) for key, var in self.controls.items()}
        try:
            output_path = self.formatter.process(gui_data, input_path)
            messagebox.showinfo("成功", f"文件已成功生成！\n保存路径: {output_path}")
        except Exception as e:
            messagebox.showerror("生成失败", f"发生严重错误：\n{e}\n\n请检查文件内容或联系技术支持。")

if __name__ == "__main__":
    app = App()
    app.mainloop()