import tkinter as tk
from tkinter import ttk, filedialog, messagebox, Toplevel, Listbox, END
import json
import os
import re
import docx
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- 全局配置 ---
CONFIG_FILE = 'format_tool_config_v4.json'

# --- 格式规范 (根据 公文要求.txt 最终版) ---
class GovDocFormatterV4:
    # 字体
    FONT_XBS = '方正小标宋简体'
    FONT_FS = '仿宋_GB2312'
    FONT_HT = '黑体'
    FONT_KT = '楷体_GB2312'
    FONT_SONG = '宋体'

    # 字号
    SIZE_2 = Pt(22)
    SIZE_3 = Pt(16)
    SIZE_4 = Pt(12)

    def set_font_style(self, run, font_name, font_size, bold=False, color_rgb=None):
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        if color_rgb:
            run.font.color.rgb = RGBColor(color_rgb[0], color_rgb[1], color_rgb[2])
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def set_paragraph_format(self, p, alignment=None, line_spacing=None, space_before=Pt(0), space_after=Pt(0), first_line_indent=None, right_indent=None):
        if alignment is not None: p.alignment = alignment
        if line_spacing is not None: 
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = line_spacing
        p.paragraph_format.space_before = space_before
        p.paragraph_format.space_after = space_after
        if first_line_indent is not None: p.paragraph_format.first_line_indent = first_line_indent
        if right_indent is not None: p.paragraph_format.right_indent = right_indent

    def add_red_separator(self, doc):
        p = doc.add_paragraph()
        self.set_paragraph_format(p, space_before=Mm(4))
        # Correct way to add a paragraph border
        pPr = p._element.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom_bdr = OxmlElement('w:bottom')
        bottom_bdr.set(qn('w:val'), 'single')
        bottom_bdr.set(qn('w:sz'), '4') # 1/8 pt, so 4 = 0.5pt
        bottom_bdr.set(qn('w:color'), 'FF0000') # Red
        p_bdr.append(bottom_bdr)
        pPr.append(p_bdr)

    def add_page_number(self, doc):
        doc.settings.odd_and_even_pages_header_footer = True
        for section in doc.sections:
            self._create_footer_page_number(section.footer, WD_ALIGN_PARAGRAPH.RIGHT)
            self._create_footer_page_number(section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT)

    def _create_footer_page_number(self, footer, alignment):
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = alignment
        p.clear()
        run = p.add_run('— ')
        self.set_font_style(run, self.FONT_SONG, self.SIZE_4)
        
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = 'PAGE'
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)

        run = p.add_run(' —')
        self.set_font_style(run, self.FONT_SONG, self.SIZE_4)

    def process(self, data, input_path):
        doc = docx.Document()
        section = doc.sections[0]
        section.top_margin = Mm(37)
        section.bottom_margin = Mm(35)
        section.left_margin = Mm(28)
        section.right_margin = Mm(26)

        if data.get('add_page_number'): self.add_page_number(doc)

        # --- 版头 ---
        if data.get('issuing_authority_logo'):
            p = doc.add_paragraph()
            self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Mm(35) - self.SIZE_2)
            run = p.add_run(data['issuing_authority_logo'])
            self.set_font_style(run, self.FONT_XBS, self.SIZE_2, color_rgb=(255, 0, 0))

        if data.get('doc_number'):
            p = doc.add_paragraph()
            self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=self.SIZE_3 * 2)
            run = p.add_run(data['doc_number'])
            self.set_font_style(run, self.FONT_FS, self.SIZE_3)

        if data.get('add_red_separator'): self.add_red_separator(doc)

        # --- 主体 ---
        if data.get('main_title'):
            p = doc.add_paragraph()
            self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=self.SIZE_3 * 2)
            run = p.add_run(data['main_title'])
            self.set_font_style(run, self.FONT_XBS, self.SIZE_2)

        if data.get('main_recipient'):
            p = doc.add_paragraph()
            self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=self.SIZE_3)
            run = p.add_run(data['main_recipient'] + '：')
            self.set_font_style(run, self.FONT_FS, self.SIZE_3)

        doc.add_paragraph() # 空一行
        self.process_body(doc, input_path, data.get('main_title'))

        # --- 文末要素 ---
        if data.get('attachment_note'):
            p = doc.add_paragraph()
            self.set_paragraph_format(p, line_spacing=Pt(28), first_line_indent=Pt(self.SIZE_3.pt * 2), space_before=self.SIZE_3)
            run = p.add_run('附件：' + data['attachment_note'])
            self.set_font_style(run, self.FONT_FS, self.SIZE_3)

        # 署名和日期 (FIXED LOGIC)
        if data.get('issuing_authority_signature') and data.get('doc_date'):
            indent_chars = 4 if data.get('is_stamped') else 2
            p = doc.add_paragraph()
            self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, 
                                      right_indent=Pt(self.SIZE_3.pt * indent_chars), 
                                      space_before=self.SIZE_3 * 2)
            
            run_sig = p.add_run(data['issuing_authority_signature'])
            self.set_font_style(run_sig, self.FONT_FS, self.SIZE_3)
            p.add_run('\n') # 换行
            run_date = p.add_run(data['doc_date'])
            self.set_font_style(run_date, self.FONT_FS, self.SIZE_3)

        output_path = os.path.splitext(input_path)[0] + "_formatted_v4.docx"
        doc.save(output_path)
        return output_path

    def process_body(self, doc, input_path, gui_title):
        is_docx = input_path.lower().endswith('.docx')
        if is_docx:
            source_doc = docx.Document(input_path)
            elements = source_doc.element.body
        else:
            with open(input_path, 'r', encoding='utf-8') as f:
                elements = f.readlines()

        for i, element in enumerate(elements):
            text_to_process = ""
            is_table = False
            if is_docx:
                if element.tag.endswith('p'):
                    text_to_process = docx.text.paragraph.Paragraph(element, doc).text.strip()
                elif element.tag.endswith('tbl'):
                    is_table = True
            else:
                text_to_process = element.strip()

            if is_table:
                table = docx.table.Table(element, doc)
                self._format_table(doc, table)
                continue

            if not text_to_process: continue
            if i == 0 and text_to_process == gui_title: continue
            self._format_paragraph(doc, text_to_process)

    def _format_paragraph(self, doc, text):
        p = doc.add_paragraph()
        self.set_paragraph_format(p, line_spacing=Pt(28))
        if re.match(r'^一、|^二、|^三、', text):
            run = p.add_run(text); self.set_font_style(run, self.FONT_HT, self.SIZE_3)
        elif re.match(r'^\uff08[一二三四五]+\uff09', text):
            run = p.add_run(text); self.set_font_style(run, self.FONT_KT, self.SIZE_3, bold=True)
        elif re.match(r'^\d+\.', text):
            run = p.add_run(text); self.set_font_style(run, self.FONT_FS, self.SIZE_3, bold=True)
        else:
            self.set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=Pt(28), first_line_indent=Pt(self.SIZE_3.pt * 2))
            run = p.add_run(text); self.set_font_style(run, self.FONT_FS, self.SIZE_3)

    def _format_table(self, doc, source_table):
        new_table = doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns), style='Table Grid')
        new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        new_table.autofit = True
        for i, row in enumerate(source_table.rows):
            for j, cell in enumerate(row.cells):
                new_table.cell(i, j).text = cell.text

# --- GUI and App Logic ---
class ConfigManager:
    def __init__(self, filename=CONFIG_FILE):
        self.filename = filename; self.data = self.load()
    def load(self):
        try:
            with open(self.filename, 'r', encoding='utf-8') as f: return json.load(f)
        except (FileNotFoundError, json.JSONDecodeError): return {}
    def save(self):
        with open(self.filename, 'w', encoding='utf-8') as f: json.dump(self.data, f, ensure_ascii=False, indent=4)
    def get(self, key, default=None): return self.data.get(key, default or [])
    def set(self, key, value): self.data[key] = value; self.save()

class ManagementDialog(Toplevel):
    def __init__(self, parent, title, key, config_manager, combobox):
        super().__init__(parent)
        self.title(f"管理 {title}"); self.key = key; self.config_manager = config_manager; self.combobox = combobox
        self.listbox = Listbox(self, width=50, height=10); self.listbox.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
        for item in self.config_manager.get(self.key): self.listbox.insert(END, item)
        btn_frame = ttk.Frame(self); btn_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        self.entry = ttk.Entry(btn_frame); self.entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(btn_frame, text="添加", command=self.add_item).pack(side=tk.LEFT, padx=(5, 0))
        ttk.Button(btn_frame, text="删除选中", command=self.delete_item).pack(side=tk.LEFT, padx=(5, 0))
    def add_item(self):
        new_item = self.entry.get()
        if new_item and new_item not in self.listbox.get(0, END):
            self.listbox.insert(END, new_item); self.entry.delete(0, END); self.save_changes()
    def delete_item(self):
        selected_indices = self.listbox.curselection()
        if not selected_indices: return
        for i in reversed(selected_indices): self.listbox.delete(i)
        self.save_changes()
    def save_changes(self):
        items = list(self.listbox.get(0, END)); self.config_manager.set(self.key, items)
        self.combobox['values'] = items
        if items: self.combobox.set(items[0])

class AppV4(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("公文智能排版工具 V4.0 - 稳定版")
        self.geometry("650x600")
        self.config_manager = ConfigManager()
        self.formatter = GovDocFormatterV4()
        self.create_widgets()

    def create_widgets(self):
        main_frame = ttk.Frame(self, padding="10"); main_frame.pack(fill=tk.BOTH, expand=True)
        file_frame = ttk.LabelFrame(main_frame, text="1. 选择文件", padding="10"); file_frame.pack(fill=tk.X)
        self.file_path_var = tk.StringVar(value="尚未选择文件...")
        ttk.Button(file_frame, text="选择 .txt 或 .docx 文件", command=self.select_file).pack(side=tk.LEFT)
        ttk.Label(file_frame, textvariable=self.file_path_var).pack(side=tk.LEFT, padx=10)
        notebook = ttk.Notebook(main_frame); notebook.pack(fill=tk.BOTH, expand=True, pady=10)
        tab1 = ttk.Frame(notebook, padding="10"); tab2 = ttk.Frame(notebook, padding="10")
        notebook.add(tab1, text='版头与主体'); notebook.add(tab2, text='文末与选项')
        self.controls = {}
        self.add_combobox(tab1, "issuing_authority_logo", "发文机关标志")
        self.add_combobox(tab1, "doc_number", "发文字号")
        self.add_entry(tab1, "main_title", "公文标题")
        self.add_combobox(tab1, "main_recipient", "主送机关") # ENHANCED
        self.add_combobox(tab2, "issuing_authority_signature", "发文机关署名")
        self.add_entry(tab2, "doc_date", "成文日期")
        self.add_entry(tab2, "attachment_note", "附件说明")
        options_frame = ttk.LabelFrame(tab2, text="格式选项", padding="10"); options_frame.pack(fill=tk.X, pady=20)
        self.controls['add_red_separator'] = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="生成红色分隔线", variable=self.controls['add_red_separator']).pack(side=tk.LEFT, padx=10)
        self.controls['is_stamped'] = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="是否加盖印章", variable=self.controls['is_stamped']).pack(side=tk.LEFT, padx=10)
        self.controls['add_page_number'] = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="生成页码", variable=self.controls['add_page_number']).pack(side=tk.LEFT, padx=10)
        ttk.Button(main_frame, text="2. 生成格式化Word文档", command=self.generate_document).pack(fill=tk.X, pady=10)

    def add_entry(self, parent, key, label):
        container = ttk.Frame(parent); container.pack(fill=tk.X, pady=(5,0))
        ttk.Label(container, text=f"{label}:").pack(anchor=tk.W)
        entry = ttk.Entry(container, width=60); entry.pack(fill=tk.X)
        self.controls[key] = entry

    def add_combobox(self, parent, key, label):
        container = ttk.Frame(parent); container.pack(fill=tk.X, pady=(5,0))
        ttk.Label(container, text=f"{label}:").pack(anchor=tk.W)
        combo_container = ttk.Frame(container)
        combo_container.pack(fill=tk.X)
        combo = ttk.Combobox(combo_container, width=58); combo.pack(side=tk.LEFT, fill=tk.X, expand=True)
        items = self.config_manager.get(key)
        combo['values'] = items
        if items: combo.set(items[0])
        self.controls[key] = combo
        ttk.Button(combo_container, text="管理...", command=lambda k=key, l=label, c=combo: self.open_management_dialog(k, l, c)).pack(side=tk.LEFT, padx=5)

    def open_management_dialog(self, key, label, combobox): ManagementDialog(self, label, key, self.config_manager, combobox)

    def select_file(self):
        path = filedialog.askopenfilename(filetypes=[("All supported", "*.txt *.docx"), ("Text", "*.txt"), ("Word", "*.docx")])
        if path: self.file_path_var.set(path); self.auto_fill_fields(path)

    def auto_fill_fields(self, path):
        text_list = []
        if path.lower().endswith('.txt'):
            with open(path, 'r', encoding='utf-8') as f: text_list = f.readlines()
        elif path.lower().endswith('.docx'):
            try:
                doc = docx.Document(path)
                text_list = [p.text for p in doc.paragraphs]
            except Exception as e: messagebox.showerror("错误", f"无法读取docx文件内容：\n{e}"); return
        if text_list:
            title = text_list[0].strip()
            if len(title) < 50: self.controls['main_title'].delete(0, END); self.controls['main_title'].insert(0, title)

    def generate_document(self):
        input_path = self.file_path_var.get()
        if "尚未选择" in input_path: messagebox.showwarning("警告", "请先选择一个文件！"); return
        gui_data = {key: (var.get() if hasattr(var, 'get') else var) for key, var in self.controls.items()}
        try:
            output_path = self.formatter.process(gui_data, input_path)
            messagebox.showinfo("成功", f"文件已成功生成！\n保存路径: {output_path}")
        except Exception as e:
            messagebox.showerror("生成失败", f"发生严重错误：\n{e}\n\n请检查文件内容或联系技术支持。")

if __name__ == "__main__":
    app = AppV4()
    app.mainloop()
